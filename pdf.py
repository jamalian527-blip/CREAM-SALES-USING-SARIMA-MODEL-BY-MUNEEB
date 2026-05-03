from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_assignment_word():
    doc = Document()

    # --- COVER PAGE ---
    # University Name
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('UNIVERSITY / INSTITUTION NAME')
    run.bold = True
    run.font.size = Pt(16)

    # Department
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Department of Islamic Studies')
    run.font.size = Pt(14)
    
    doc.add_paragraph() # Spacer

    # Course Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('COURSE TITLE: ISLAMIC STUDIES')
    run.bold = True
    run.font.size = Pt(14)

    doc.add_paragraph() # Spacer

    # Topic
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Assignment Topic: Can a society maintain moral and social health if the family unit is weak or fragmented?')
    run.bold = True
    run.font.size = Pt(12)

    doc.add_paragraph() # Spacer
    doc.add_paragraph() # Spacer

    # Student Details
    details = [
        "Instructor Name: [Insert Instructor Name]",
        "Student Name: [Insert Your Name]",
        "Roll No: [Insert Roll No]",
        "Semester/Section: [Insert Semester]",
        "Submission Date: [Insert Date]"
    ]

    for line in details:
        p = doc.add_paragraph(line)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.runs[0].font.size = Pt(12)

    doc.add_page_break()

    # --- BODY CONTENT ---
    
    def add_heading(text):
        h = doc.add_heading(text, level=1)
        h.runs[0].font.name = 'Arial'
        h.runs[0].font.color.rgb = None # Standard black

    def add_body_text(text):
        p = doc.add_paragraph(text)
        p.runs[0].font.name = 'Times New Roman'
        p.runs[0].font.size = Pt(12)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Content
    add_heading("1. Introduction")
    add_body_text("In the intricate architecture of human civilization, the family stands as the foundational cornerstone. From the Islamic perspective, the family is not merely a social construct or a biological necessity; it is a divine institution sanctified by Allah (SWT) and established by the Prophets. The topic under discussion, \"Can a society maintain moral and social health if the family unit is weak or fragmented?\", strikes at the very heart of Islamic sociology.")
    add_body_text("The concept of the Ummah (the global Muslim community) is often visualized as a single body. However, the cells that make up this body are the individual family units. Islam posits that the strength, morality, and stability of the collective society are entirely dependent on the spiritual and functional health of the household. When the family creates an environment of piety, discipline, and love, the society flourishes. Conversely, when the family unit is weak, fragmented, or neglected, the social fabric inevitably unravels, leading to moral decay and anarchy.")
    add_body_text("This assignment explores the critical role of the family as the \"building block of the Ummah.\" It argues that maintaining moral and social health is impossible in the absence of strong families.")

    add_heading("2. Conceptual & Theoretical Framework")
    add_body_text("To understand why the family is indispensable to social health, we must first establish the Islamic theoretical framework regarding the \"family\" (Usrah).")
    add_body_text("The Family as a Microcosm of Society: In Islamic theory, the family is the microcosm of the macrocosm (society). It is the first training ground—the primary school of civilization—where a human being learns the essential values of interaction: authority, compassion, rights, and responsibilities. Imam Al-Ghazali theorized that the governance of the state is merely an extension of the governance of the home. If there is tyranny or chaos in the home, there will be tyranny and chaos in the state.")
    add_body_text("Key Islamic Concepts:\n1. Divine Sanctity (Mithaq Ghaliz): The Qur’an refers to the marital bond as a solemn covenant.\n2. Lineage (Nasab): Islam places immense weight on the preservation of genealogy.\n3. Tarbiyah (Upbringing): The family is the vessel through which the Deen is transmitted.")

    add_heading("3. Qur'anic Perspective")
    add_body_text("The Holy Qur'an provides the constitution for the family, emphasizing its role as a source of peace and a preventative measure against social evil.")
    add_body_text("1. The Source of Social Tranquility (Sakinah) - Surah Ar-Rum (30:21): \"And of His signs is that He created for you from yourselves mates that you may find tranquillity in them...\" Context: Sakinah implies that the family is the cure for anxiety.")
    add_body_text("2. The Mechanism of Moral Protection - Surah At-Tahrim (66:6): \"O you who have believed, protect yourselves and your families from a Fire...\" Context: This establishes the family as a defensive unit against immorality.")

    add_heading("4. Prophetic Guidance - Ahadith")
    add_body_text("The Sunnah of the Prophet Muhammad (PBUH) translates the Qur'anic theory into practical application.")
    add_body_text("1. Accountability (Sahih Al-Bukhari): \"All of you are shepherds...\" Implication: Social order is decentralized to the family unit.")
    add_body_text("2. Excellence (Sunan At-Tirmidhi): \"The best of you is the one who is best to his family...\" Implication: Private conduct determines public social health.")
    
    add_heading("5. Contemporary Relevance / Application")
    add_body_text("In the 21st century, we are witnessing a global experiment in what happens when the family unit is de-prioritized. Modern society is grappling with an epidemic of mental health issues and juvenile delinquency. From an Islamic lens, these are symptoms of the collapse of the home.")
    add_body_text("Islamic Solutions: Islam reorients the definition of \"success\" to prioritize family stability (Falah) over pure economic ambition, ensuring the next generation is psychologically secure.")

    add_heading("6. Critical Analysis / Reflection")
    add_body_text("Can't a strong state or strong education system compensate for weak families? The answer is no. Institutions are reactive; families are proactive. When the family fragments, the state is forced to intrude more into private life (e.g., policing, welfare). The Islamic model aims for a self-regulating society where strong families manage their own affairs and morals.")

    add_heading("7. Conclusion")
    add_body_text("In conclusion, a society cannot maintain moral and social health if the family unit is weak. The family is the bedrock of the Ummah. The fragmentation of the family unit is the precursor to the fragmentation of society itself. To heal the Ummah, we must first heal the home.")

    add_heading("References / Bibliography")
    add_body_text("1. Al-Qur'an Al-Kareem (Surah Ar-Rum, At-Tahrim, An-Nisa)\n2. Sahih Al-Bukhari (Book of Nikah)\n3. Sahih Muslim (Book of Birr wa'l-Silah)\n4. Jami` at-Tirmidhi (Chapters on Virtues)\n5. Ulwan, Abdullah Nasih. Tarbiyat al-Awlad fi al-Islam")

    doc.save('Islamic_Studies_Assignment.docx')
    print("Word file created: Islamic_Studies_Assignment.docx")

if __name__ == '__main__':
    create_assignment_word()